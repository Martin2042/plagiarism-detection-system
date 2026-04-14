import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

######################################

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

def get_url(filename):
    urls_list = []
    docx_file=filename
    document = Document(docx_file)
    rels = document.part.rels
    sections = document.sections
    for section in sections:
        print(section.start_type)

    for rel in rels:
        #print(rels[rel].reltype)
        if rels[rel].reltype == RT.HYPERLINK:
            urls_list.append(rels[rel]._target)
            print("\n Origianl link id -", rel, "with detected URL: ", rels[rel]._target)

    return urls_list

if __name__ == '__main__':
    st = getText('F:/workspace/PycharmProjects/MES_MCA_2022_PlagarismDetection/test.docx')
    print(st.strip())

    urls_list = get_url('F:/workspace/PycharmProjects/MES_MCA_2022_PlagarismDetection/test.docx')
    print(urls_list)
